from io import BytesIO
import locale
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from PIL import Image

# ضع المتغيرات العالمية هنا 
existing_docx_file = './daily_diary.docx'  # يجب ان اغير مكان الملف و اجعل الكود يتعرف على مكانه تلقائيا
MAX_SIZE = 350

# ضع الكود الجديد هنا  في البداية

def resize_image(img):
    width, height = img.size
    if width > MAX_SIZE or height > MAX_SIZE:
        if width > height:
            new_width = MAX_SIZE
            new_height = int(height * (MAX_SIZE / width))
        else:
            new_height = MAX_SIZE
            new_width = int(width * (MAX_SIZE / height))
            img = img.resize((new_width, new_height), resample=Image.BILINEAR)
    return img

def insert_image_to_docx(image_file, file_path):
    # Create a new Document
    doc = Document(file_path)

    # Create a new PIL Image from the image byte stream
    with BytesIO() as image_buffer:
        image_file.download(out=image_buffer)
        image_buffer.seek(0)
        img = Image.open(image_buffer)
        
        # Resize the image using the resize_image function
        img = resize_image(img)

        # Add the resized image to the DOCX file
        img_bytes = BytesIO()
        img.save(img_bytes, format='JPEG')
        # doc.add_picture(img_bytes)
        
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(img_bytes)
        # Align the paragraph to the right
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_paragraph(style='LO-normal' )
        
        # Save the document
        doc.save(file_path)
        print(f"picture added to {file_path}")
        
        # # Save the DOCX document with a unique filename
        # doc.save(docx_filename)
  
def get_the_current_time():
    current_time = datetime.now()
    time_str = current_time.strftime("%I:%M %p").replace('PM', 'م').replace('AM', 'ص')
    return time_str

def format_arabic_date_with_day_name():
    # Set the locale to Arabic
    locale.setlocale(locale.LC_TIME, 'ar_AE.UTF-8')  # Change 'ar_AE' to your desired Arabic locale

    # Get the current date and time
    current_date = datetime.now()

    # Format the date to include the day name in Arabic
    formatted_date = current_date.strftime("%A، %Y/%m/%d")

    return formatted_date

def add_content_to_existing_docx(file_path , text ,style='LO-normal' ,write_time_heading=False):
    # يوجد نوعان للمتغير style
    # LO-normal , LO-heading
    # Open the existing Document
    doc = Document(file_path)

    if write_time_heading:
        doc.add_paragraph(get_the_current_time() , style='LO-heading' )
        
    # Add paragraph with style
    doc.add_paragraph(text , style=style )
    doc.add_paragraph(style='LO-normal' )

    # Save the modified document back to the same file
    doc.save(file_path)

    print(f"Content added to {file_path}")

# Call the function to add content to the existing document
def add_date_to_existing_doc(existing_docx_file ):
    arabic_date = format_arabic_date_with_day_name()
    add_content_to_existing_docx(existing_docx_file , arabic_date ,style='LO-heading')

if __name__ == "__main__":
    # existing_docx_file = './daily_diary.docx'
    get_the_current_time()
    # add_date_to_existing_doc('./daily_diary.docx')
    insert_image_to_docx('./image.png', './daily_diary.docx')